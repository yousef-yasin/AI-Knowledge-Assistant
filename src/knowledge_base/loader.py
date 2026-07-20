from __future__ import annotations # Enable postponed evaluation of type annotations. 

import csv # Read CSV files.
from dataclasses import dataclass   # Create simple data container classes.
from pathlib import Path  # Work with file system paths.
from typing import Any # Allow metadata values of any type.

from docx import Document as DocxDocument # Read Microsoft Word documents.
from pypdf import PdfReader  # Read PDF files.

# Import helper functions for creating document metadata.
from src.knowledge_base.metadata import add_page_metadata, create_base_metadata

# Store the extracted document text together with its metadata.
@dataclass
class LoadedDocument: 
    text: str # Extracted document text.
    metadata: dict[str, Any] # Information such as filename and page number.

    # Supported file extensions.
class DocumentLoader:
    """Load PDF, DOCX, TXT, Markdown and CSV documents, with OCR fallback."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}

    def __init__(self, enable_ocr: bool = True, ocr_language: str = "eng") -> None:
        self.enable_ocr = enable_ocr     # Enable OCR for scanned PDF pages.

        self.ocr_language = ocr_language      # OCR language used by Tesseract.


    def load(self, file_path: str | Path) -> list[LoadedDocument]:
    # Convert the input path into a Path object.

        path = Path(file_path)
        # Ensure the file exists.
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        # Ensure the path points to a file.
        if not path.is_file():
            raise ValueError(f"The provided path is not a file: {path}")
        # Get the file extension.
        extension = path.suffix.lower()
        # Validate that the file type is supported.
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        # Validate that the file type is supported.
        loaders = {
            ".pdf": self._load_pdf, ".docx": self._load_docx,
            ".txt": self._load_text, ".md": self._load_text,
            ".csv": self._load_csv,
        }
        # Load the document and ignore empty results.
        documents = [doc for doc in loaders[extension](path) if doc.text.strip()]
        # Raise an error if no readable content was found.
        if not documents:
            
            raise ValueError(f"No readable text was found in: {path.name}")
        return documents
# Read every page in the PDF.
    def _load_pdf(self, path: Path) -> list[LoadedDocument]:
        reader = PdfReader(str(path))
        # Store all extracted pages.
        documents: list[LoadedDocument] = []
        # Process each page in the PDF.
        for page_index, page in enumerate(reader.pages, start=1):
            # Try to extract text directly from the page.
            text = (page.extract_text() or "").strip()
            # Assume the text came from native PDF extraction.
            extraction_method = "native_text"
            # Use OCR if the extracted text is too short.
            if self.enable_ocr and len(text) < 20:
                ocr_text = self._ocr_pdf_page(path, page_index - 1)
                if ocr_text.strip():
                    # Replace the extracted text if OCR succeeded.
                    text = ocr_text.strip()
                    extraction_method = "ocr"
            # Create metadata for the current page.
            metadata = add_page_metadata(
                create_base_metadata(path), page_index, len(reader.pages)
            )
            # Record how the text was extracted.
            metadata["extraction_method"] = extraction_method
            # Save the extracted page.
            documents.append(LoadedDocument(text=text, metadata=metadata))
        # Return all extracted pages.
        return documents

    def _ocr_pdf_page(self, path: Path, page_index: int) -> str:
        """Render one PDF page and OCR it. Returns empty text if OCR is unavailable."""
        try:
            # Import OCR dependencies only when needed.
            import pypdfium2 as pdfium
            import pytesseract
            # Open the PDF document.
            pdf = pdfium.PdfDocument(str(path))
            # Select the requested page.
            page = pdf[page_index]
            # Convert the page into an image.
            image = page.render(scale=2.5).to_pil()
            return pytesseract.image_to_string(image, lang=self.ocr_language)
        # Return an empty string if OCR fails.
        except Exception:
            return ""

    def _load_docx(self, path: Path) -> list[LoadedDocument]:
        # Open the Word document.
        document = DocxDocument(str(path))
        # Collect all non-empty paragraphs.
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        # Create metadata for the document.
        metadata = add_page_metadata(create_base_metadata(path), None)
        # Return the document as a single text block.
        return [LoadedDocument("\n".join(paragraphs), metadata)]

    def _load_text(self, path: Path) -> list[LoadedDocument]:
        # Create metadata for the text file.
        metadata = add_page_metadata(create_base_metadata(path), None)
        return [LoadedDocument(self._read_text_with_fallback(path).strip(), metadata)]

    def _load_csv(self, path: Path) -> list[LoadedDocument]:
        # Store all formatted rows.
        rows: list[str] = []
        # Open the CSV file.
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            # Read the CSV as dictionaries.
            reader = csv.DictReader(handle)
            # Ensure the file contains column headers.
            if reader.fieldnames is None:
                raise ValueError(f"CSV file has no headers: {path.name}")
            # Process each row.
            for number, row in enumerate(reader, 1):
                values = [f"{key}: {value}" for key, value in row.items()
                          if value is not None and str(value).strip()]
                if values: # Keep only non-empty rows.
                    rows.append(f"Row {number}: " + " | ".join(values))
        # Create metadata for the CSV file.
        metadata = add_page_metadata(create_base_metadata(path), None)
        # Store the total number of rows.
        metadata["row_count"] = len(rows)
        # Return all rows as one document.
        return [LoadedDocument("\n".join(rows), metadata)]

    @staticmethod
    def _read_text_with_fallback(path: Path) -> str:
        # Try reading the file using common text encodings.
        for encoding in ("utf-8", "utf-8-sig", "utf-16", "cp1252"):
            try: # Return the file content if decoding succeeds.
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:  # Try the next encoding if decoding fails.
                continue
        # Raise an error if all decoding attempts fail.
        raise ValueError(f"Could not decode text file: {path.name}")
